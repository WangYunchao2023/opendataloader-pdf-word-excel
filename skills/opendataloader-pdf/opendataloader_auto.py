#!/usr/bin/env python3
"""
文档自动检测与转换脚本（统一处理 PDF / Word / Excel）
版本: 1.6.0
日期: 2026-04-01

架构说明（v1.6）：
  输入为 PDF  → 直接提取（位置信息 + 内容，PDF 限制：有 bbox/页码，表格数据可能稀疏）
  输入为 Word → Word转PDF + docx直接提取 + PDF位置提取 → 智能合并
               相同信息以更可靠来源为准：表格/文本内容以 docx 为准，位置信息（页码/bbox）以 PDF 为准
  输入为 Excel → openpyxl直接提取（结构化数据/表格/图表）+ Excel转PDF获取页码 → 合并

输出：统一的两文件架构
  ✦ {basename}.json    ← 结构化数据，含完整内容和位置信息
  ✦ {basename}.md      ← 人类可读的 Markdown
"""

import sys
import os
import json
import argparse
import subprocess
import time
import re
from pathlib import Path

# ---------- Java 路径 ----------
JAVA_BIN = "/home/wangyc/opt/jre/amazon-corretto-11.0.30.7.1-linux-x64/bin/java"
JAVA_HOME = "/home/wangyc/opt/jre/amazon-corretto-11.0.30.7.1-linux-x64"

HYBRID_SERVER_PID_FILE = "/tmp/opendataloader_hybrid_server.pid"
HYBRID_SERVER_PORT = 5002

# ---------- Word 转 PDF ----------
def convert_word_to_pdf(docx_path: str, output_dir: str = "/tmp") -> str:
    """
    用 LibreOffice 将 Word 文档转换为 PDF
    返回: PDF 文件的绝对路径
    """
    import tempfile

    pdf_dir = output_dir if output_dir else "/tmp"
    os.makedirs(pdf_dir, exist_ok=True)

    basename = Path(docx_path).stem + ".pdf"
    pdf_path = os.path.join(pdf_dir, basename)

    cmd = [
        "libreoffice", "--headless",
        "--convert-to", "pdf",
        "--outdir", pdf_dir,
        os.path.abspath(docx_path),
    ]

    print(f"[INFO] Word → PDF 转换中: {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice 转换失败: {result.stderr[:300]}")

    # LibreOffice 可能改名（如加编号），找最新生成的 pdf
    # 等待文件落盘
    time.sleep(1)

    # 精确路径：与 docx 同名的 pdf
    expected = os.path.join(pdf_dir, basename)
    if os.path.exists(expected):
        print(f"[INFO] PDF 生成成功: {expected}")
        return expected

    # 备选：找最新的 pdf 文件
    files = sorted(Path(pdf_dir).glob("*.pdf"), key=lambda p: p.stat().st_mtime, reverse=True)
    if files:
        print(f"[INFO] PDF 生成成功: {files[0]}")
        return str(files[0])

    raise RuntimeError("PDF 转换后未找到输出文件")


# ---------- PDF 提取结果增强（内容指纹） ----------
def add_content_fingerprint(data: dict) -> dict:
    """
    给 PDF 提取的 JSON 添加内容指纹，增强可追溯性。
    1. 扁平化 kids 嵌套结构 → flat_elements
    2. 添加 section_path（标题层级路径）
    3. 添加 table_index / paragraph_index（全局序号）
    4. 添加 content_preview（前100字符，用于快速定位）
    """
    flat = []
    heading_path = []  # 栈：当前标题路径

    def flatten(kids, flat_list, heading_path):
        for kid in kids:
            elem = kid.copy()
            kid_type = elem.get("type", "")

            # 维护标题路径
            if kid_type in ("heading", "header"):
                level = elem.get("heading level", 99)
                # Doctitle=0, Subtitle=1, 正文=2...
                # 简单映射：heading level 数值
                # 弹出更高级的标题
                while heading_path and heading_path[-1][0] >= level:
                    heading_path.pop()
                heading_path.append((level, elem.get("content", "")))
                elem["section_path"] = " > ".join(h for _, h in heading_path)
                flat_list.append(elem)
                # 递归处理 kids（子元素）
                if "kids" in elem:
                    flatten(elem.pop("kids"), flat_list, heading_path)
            else:
                flat_list.append(elem)
                if "kids" in elem:
                    flatten(elem.pop("kids"), flat_list, heading_path)

    if "kids" in data:
        flatten(data["kids"], flat, heading_path)

    # 全局序号
    table_count = 0
    para_count = 0
    img_count = 0

    for elem in flat:
        t = elem.get("type", "")
        content = elem.get("content", "")

        # 内容预览（前100字）
        elem["content_preview"] = content[:100] + ("..." if len(content) > 100 else "")

        if t == "table":
            table_count += 1
            elem["table_index"] = table_count
        elif t == "paragraph":
            para_count += 1
            elem["paragraph_index"] = para_count
        elif t in ("image", "figure"):
            img_count += 1
            elem["image_index"] = img_count

    # 元数据
    data["traceability"] = {
        "total_elements": len(flat),
        "total_tables": table_count,
        "total_paragraphs": para_count,
        "total_images": img_count,
        "fingerprint_version": "1.0",
    }
    data["flat_elements"] = flat
    return data


# ---------- 核心：合并 Word 内容 + PDF 位置信息 ----------
def normalize_text(text: str) -> str:
    """文本归一化：去空格、标点、换行，用于匹配"""
    if not text:
        return ""
    import unicodedata
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"[\s\n\r\t\xa0·\-–—–‑]", "", text)
    text = text.lower()
    return text


def text_similarity(t1: str, t2: str) -> float:
    """两个文本的相似度（0~1）"""
    n1, n2 = normalize_text(t1), normalize_text(t2)
    if not n1 or not n2:
        return 0.0
    if n1 == n2:
        return 1.0
    # 简单：n1 在 n2 中或 n2 在 n1 中 → 高相似
    if n1 in n2 or n2 in n1:
        return 0.8
    # Jaccard
    s1, s2 = set(n1), set(n2)
    if not s1 or not s2:
        return 0.0
    return len(s1 & s2) / len(s1 | s2)


def extract_pdf_positions(pdf_flat: list) -> dict:
    """
    从 PDF 扁平元素列表中提取：
    - 标题层级路径（heading → section_path）
    - 页码 + bbox（heading / table / paragraph / image）
    返回结构：{(type, normalized_key): {page, bbox, section_path}}
    """
    pos_map = {}
    heading_path = []

    for elem in pdf_flat:
        t = elem.get("type", "")
        content = elem.get("content", "")

        if t == "heading":
            level = elem.get("heading level", 99)
            heading_path.append((level, content))
            heading_path.sort(key=lambda x: x[0])
            path_str = " > ".join(h for _, h in heading_path)
            key = normalize_text(content)
            pos_map[("heading", key)] = {
                "page": elem.get("page number"),
                "bbox": elem.get("bounding box"),
                "section_path": path_str,
            }
        elif t in ("table", "paragraph", "image", "figure"):
            key = normalize_text(content)[:50]
            if ("table", key) not in pos_map:
                pos_map[(t, key)] = {
                    "page": elem.get("page number"),
                    "bbox": elem.get("bounding box"),
                }

    return pos_map


def find_best_pdf_match(docx_elem: dict, pdf_pos_map: dict,
                        seen_tables: set, seen_paras: set) -> dict:
    """
    为单个 docx 元素在 pdf_flat 中找最佳匹配的 position 信息。
    匹配策略：
      - 标题：按 text 精确匹配（normalize 后）
      - 表格：按内容相似度匹配（取第一行文字）
      - 段落：按内容相似度匹配（跳过已匹配的）
    返回：{page, bbox, section_path} 或 {}
    """
    t = docx_elem.get("type", "")
    content = docx_elem.get("content", "")

    norm_key = normalize_text(content)

    if t == "heading":
        for key, info in pdf_pos_map.items():
            if key[0] == "heading" and norm_key and key[1] == norm_key:
                return {"page": info["page"], "bbox": info["bbox"],
                        "section_path": info["section_path"]}
        # 模糊匹配
        for key, info in pdf_pos_map.items():
            if key[0] == "heading" and norm_key and text_similarity(content, key[1]) > 0.7:
                return {"page": info["page"], "bbox": info["bbox"],
                        "section_path": info["section_path"]}

    elif t == "table":
        # 取表格第一行做匹配键
        lines = content.strip().split("\n")
        if len(lines) >= 2:
            header = lines[1] if lines[0].startswith("|") else lines[0]
            norm_header = normalize_text(header)[:50]
            for key, info in pdf_pos_map.items():
                if key[0] == "table" and norm_header and key[1] == norm_header:
                    if key not in seen_tables:
                        seen_tables.add(key)
                        return {"page": info["page"], "bbox": info["bbox"]}
        # 模糊匹配
        for key, info in pdf_pos_map.items():
            if key[0] == "table" and key not in seen_tables:
                seen_tables.add(key)
                return {"page": info["page"], "bbox": info["bbox"]}

    elif t == "paragraph":
        if not norm_key:
            return {}
        for key, info in pdf_pos_map.items():
            if key[0] == "paragraph" and key not in seen_paras:
                sim = text_similarity(content, key[1])
                if sim > 0.7:
                    seen_paras.add(key)
                    return {"page": info["page"], "bbox": info["bbox"]}

    return {}


def merge_docx_and_pdf(docx_data: dict, pdf_flat: list) -> dict:
    """
    核心合并函数：以 docx 内容为主体，为每个元素附加 PDF 位置信息。
    相同信息以更可靠来源为准：表格/文本内容以 docx 为准，页码/bbox 以 PDF 为准。
    """
    # 构建 PDF position 字典
    pdf_pos_map = extract_pdf_positions(pdf_flat)
    seen_tables = set()
    seen_paras = set()

    # 全局序号
    table_count = 0
    para_count = 0
    img_count = 0

    # 收集 section_path（从 docx 的标题层级）
    heading_path = []

    merged_elements = []
    for elem in docx_data.get("elements", []):
        elem = dict(elem)  # 浅拷贝
        t = elem.get("type", "")

        # 更新标题路径
        if t == "heading":
            level = elem.get("heading_level", 99)
            content = elem.get("content", "")
            while heading_path and heading_path[-1][0] >= level:
                heading_path.pop()
            heading_path.append((level, content))
            elem["section_path"] = " > ".join(h for _, h in heading_path)

        # 内容预览
        elem["content_preview"] = elem.get("content", "")[:100]
        if len(elem.get("content", "")) > 100:
            elem["content_preview"] += "..."

        # 在 PDF 中找位置（最佳匹配）
        pos = find_best_pdf_match(elem, pdf_pos_map, seen_tables, seen_paras)
        if pos:
            elem["page"] = pos.get("page")
            elem["bbox"] = pos.get("bbox")

        # 序号
        if t == "table":
            table_count += 1
            elem["table_index"] = table_count
        elif t == "paragraph":
            para_count += 1
            elem["paragraph_index"] = para_count

        merged_elements.append(elem)

    # 元数据
    result = {
        "doc_type": "word",
        "source": "word-to-pdf-merged",
        "total_elements": len(merged_elements),
        "total_tables": table_count,
        "total_paragraphs": para_count,
        "elements": merged_elements,
        # PDF 位置信息（备用）
        "_pdf_position_hints": {
            "total_from_pdf": len(pdf_flat),
            "tables_matched": len(seen_tables),
            "paras_matched": len(seen_paras),
        }
    }
    return result


def convert_word_to_markdown(docx_path: str) -> str:
    """生成可读 Markdown（基于 docx，与合并后 JSON 对应）"""
    from docx import Document
    doc = Document(docx_path)
    lines = []
    heading_path = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading") or (para.style and "Heading" in para.style.name):
            level = 1
            for c in style_name:
                if c.isdigit():
                    level = int(c)
                    break
            lines.append(f"{'#' * level} {text}")
            # 维护路径
            while heading_path and heading_path[-1][0] >= level:
                heading_path.pop()
            heading_path.append((level, text))
        else:
            lines.append(text)

    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            rows_text.append(row_text)
        if rows_text:
            lines.append("")
            lines.append(rows_text[0])
            lines.append(" | ".join(["---"] * len(rows_text[0].split(" | "))))
            for row in rows_text[1:]:
                lines.append(row)

    return "\n".join(lines)


# ---------- Excel 提取 ----------
def _make_serializable(val):
    """将 Excel 单元格值转为 JSON 可序列化类型"""
    import datetime
    if val is None:
        return ""
    if isinstance(val, (datetime.datetime, datetime.date, datetime.time)):
        return val.isoformat()
    if isinstance(val, float):
        # 截断浮点数精度
        return round(val, 6)
    return val


def extract_excel_to_json(xlsx_path: str) -> dict:
    """
    用 openpyxl 提取 Excel 完整内容：
    - 工作表列表 + 表头
    - 每个 sheet 的数据区域（转为 table 结构）
    - 命名区域 / 表格对象
    - chart 图表（提取图表类型 + 关联数据表）
    输出格式与 PDF/Word 统一（elements 数组）。
    """
    import openpyxl
    from openpyxl.utils import get_column_letter
    from openpyxl.chart import BarChart, LineChart, PieChart, ScatterChart, Reference

    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    elements = []
    sheet_count = 0
    table_count = 0
    chart_count = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_count += 1

        # 跳过空 sheet
        if ws.max_row < 2 or ws.max_column < 2:
            continue

        # 判断 sheet 是否为表格（是否有数据的起始行）
        # 收集该 sheet 的所有非空单元格数据
        data_rows = []
        for row in ws.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                data_rows.append(list(row))

        if not data_rows:
            continue

        # 检测表头（第一行是否可作为列名）
        headers = [_make_serializable(c) for c in data_rows[0]]

        # 构建 Markdown 表格
        md_rows = []
        md_rows.append("| " + " | ".join(str(h) for h in headers) + " |")
        md_rows.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in data_rows[1:]:
            row_vals = [_make_serializable(c) for c in row]
            row_str = " | ".join(str(v) for v in row_vals)
            md_rows.append("| " + row_str + " |")

        table_content = "\n".join(md_rows)

        # 可序列化的 data_rows
        serializable_rows = [
            [_make_serializable(c) for c in row] for row in data_rows
        ]

        elements.append({
            "type": "table",
            "location": {
                "excel": sheet_name,           # 主要标识：原Excel Sheet名
                "pdf_page": None,               # PDF页码（匹配到则填入，否则None）
            },
            "table_index": table_count + 1,
            "content": table_content,
            "content_preview": table_content[:100],
            "row_count": len(data_rows),
            "col_count": len(headers),
            "headers": headers,
            "data_rows": serializable_rows,     # 结构化数组，AI 可直接分析
            "section_path": f"Excel > {Path(xlsx_path).stem} > {sheet_name}",
        })
        table_count += 1

        # 提取 chart 信息
        if hasattr(ws, "_charts"):
            for chart in ws._charts:
                chart_count += 1
                chart_type = type(chart).__name__
                chart_title = chart.title if hasattr(chart, "title") and chart.title else ""

                # 提取图表关联的数据范围
                data_desc = ""
                if hasattr(chart, "data"):
                    refs = chart.data
                    if refs:
                        data_desc = f"数据范围: {refs}"

                # 构建图表的摘要数据（将图表还原为简化的数据表）
                sample_data = []
                if hasattr(chart, "series") and chart.series:
                    for series in chart.series:
                        sname = series.title if hasattr(series, "title") and series.title else f"系列{chart_count}"
                        sample_data.append(f"系列: {sname}")

                elements.append({
                    "type": "chart",
                    "location": {
                        "excel": sheet_name,           # 主要标识：原Excel Sheet名
                        "pdf_page": None,               # PDF页码（匹配到则填入，否则None）
                    },
                    "chart_index": chart_count,
                    "chart_type": chart_type,
                    "chart_title": str(chart_title),
                    "content": f"[{chart_type}] {chart_title} | {data_desc}",
                    "content_preview": f"{chart_type}: {chart_title}",
                    "section_path": f"Excel > {Path(xlsx_path).stem} > {sheet_name} > 图表",
                })

    # 命名区域
    defined_names = []
    if wb.defined_names:
        for name, defn in wb.defined_names.items():
            defined_names.append({
                "name": name,
                "value": _make_serializable(getattr(defn, "value", "")),
                "attr_text": str(getattr(defn, "attr_text", "")),
            })

    return {
        "doc_type": "excel",
        "source": "openpyxl",
        "total_sheets": sheet_count,
        "total_tables": table_count,
        "total_charts": chart_count,
        "defined_names": defined_names,
        "sheets": wb.sheetnames,
        "elements": elements,
    }


def extract_excel_to_markdown(xlsx_path: str, page_hints: dict = None) -> str:
    """生成 Excel 的可读 Markdown（每个 sheet 转为 Markdown 表格）"""
    import openpyxl

    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    page_hints = page_hints or {}
    lines = [f"# {Path(xlsx_path).stem}\n"]

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        page_hint = page_hints.get(sheet_name, "")
        pg_note = f" _(PDF 第{page_hint}页)_" if page_hint else ""
        lines.append(f"\n## Sheet: {sheet_name}{pg_note}\n")

        data_rows = []
        for row in ws.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                data_rows.append(list(row))

        if not data_rows:
            lines.append("_（空表）_\n")
            continue

        headers = [_make_serializable(c) for c in data_rows[0]]
        lines.append("| " + " | ".join(str(h) for h in headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in data_rows[1:]:
            row_vals = [_make_serializable(c) for c in row]
            row_str = " | ".join(str(v) for v in row_vals)
            lines.append("| " + row_str + " |")
        lines.append("")

    return "\n".join(lines)


# ---------- Word 提取 ----------
    """用 python-docx 提取 Word 内容，输出与 PDF 相同的 JSON 结构"""
    from docx import Document

    doc = Document(docx_path)
    elements = []
    page = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 判断段落类型
        style_name = para.style.name.lower() if para.style else ""

        if style_name.startswith("heading") or para.style and "Heading" in para.style.name:
            # 提取标题级别
            level = 1
            for c in para.style.name:
                if c.isdigit():
                    level = int(c)
                    break
            elements.append({
                "type": "heading",
                "content": text,
                "heading_level": level,
                "page": page,
                "bbox": None,
            })
        else:
            elements.append({
                "type": "paragraph",
                "content": text,
                "page": page,
                "bbox": None,
            })

    # 提取表格
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_texts)

        if rows_data:
            # 构建 Markdown 表格
            md_table = ""
            if rows_data:
                # 表头
                md_table += "| " + " | ".join(rows_data[0]) + " |\n"
                # 分隔线
                md_table += "| " + " | ".join(["---"] * len(rows_data[0])) + " |\n"
                # 数据行
                for row in rows_data[1:]:
                    md_table += "| " + " | ".join(row) + " |\n"

            elements.append({
                "type": "table",
                "content": md_table.strip(),
                "page": page,
                "bbox": None,
            })

    # 转为与 PDF 输出兼容的 JSON
    return {
        "doc_type": "word",
        "total_pages": 1,
        "elements": elements,
    }


def extract_word_to_markdown(docx_path: str) -> str:
    """用 python-docx 提取 Word 内容，输出 Markdown（与 PDF 提取的 .md 格式一致）"""
    from docx import Document

    doc = Document(docx_path)
    lines = []
    in_table = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if in_table:
                lines.append("")
                in_table = False
            continue

        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading") or (para.style and "Heading" in para.style.name):
            # 标题
            level = 1
            for c in style_name:
                if c.isdigit():
                    level = int(c)
                    break
            lines.append(f"{'#' * level} {text}")
            in_table = False
        else:
            lines.append(text)
            in_table = False

    # 处理表格
    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            rows_text.append(row_text)

        if rows_text:
            lines.append("")
            lines.append(rows_text[0])  # 表头
            lines.append(" | ".join(["---"] * len(rows_text[0].split(" | "))))  # 分隔线
            for row in rows_text[1:]:
                lines.append(row)

    return "\n".join(lines)


def convert_word(input_path: str, output_dir: str, output_format: str = "markdown,json") -> dict:
    """转换 Word 文档，统一输出格式"""
    os.makedirs(output_dir, exist_ok=True)

    basename = Path(input_path).stem
    formats = output_format.split(",")

    result = {
        "success": True,
        "input_path": input_path,
        "output_dir": output_dir,
        "files_created": [],
        "mode": "word-direct",
    }

    if "json" in formats:
        json_data = extract_word_to_json(input_path)
        json_path = os.path.join(output_dir, f"{basename}.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        result["files_created"].append(json_path)

    if "markdown" in formats:
        md_data = extract_word_to_markdown(input_path)
        md_path = os.path.join(output_dir, f"{basename}.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(f"# {basename}\n\n")
            f.write(md_data)
        result["files_created"].append(md_path)

    if "text" in formats:
        from docx import Document
        doc = Document(input_path)
        text_path = os.path.join(output_dir, f"{basename}.txt")
        with open(text_path, "w", encoding="utf-8") as f:
            for para in doc.paragraphs:
                f.write(para.text + "\n")
        result["files_created"].append(text_path)

    # 如果格式都不对，默认输出 markdown 和 json
    if not result["files_created"]:
        json_data = extract_word_to_json(input_path)
        md_data = extract_word_to_markdown(input_path)

        json_path = os.path.join(output_dir, f"{basename}.json")
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        result["files_created"].append(json_path)

        md_path = os.path.join(output_dir, f"{basename}.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(f"# {basename}\n\n")
            f.write(md_data)
        result["files_created"].append(md_path)

    return result


def find_java():
    """检查 Java 是否可用"""
    try:
        result = subprocess.run(
            [JAVA_BIN, "-version"],
            capture_output=True, text=True, timeout=5
        )
        return True
    except Exception:
        return False


def detect_pdf_type(pdf_path: str) -> dict:
    """
    检测 PDF 类型
    返回: {
        "type": "digital" | "scanned" | "mixed",
        "lang": "auto" | "zh" | "en" | "mixed",
        "has_complex_tables": bool,
        "has_formulas": bool,
        "has_charts": bool,
        "total_pages": int,
        "text_pages": int,
        "scanned_pages": int,
    }
    """
    result = {
        "type": "digital",
        "lang": "auto",
        "has_complex_tables": False,
        "has_formulas": False,
        "has_charts": False,
        "total_pages": 0,
        "text_pages": 0,
        "scanned_pages": 0,
    }

    try:
        from pypdf import PdfReader
        reader = PdfReader(pdf_path)
        result["total_pages"] = len(reader.pages)

        # 检查前5页（统计有文字的页数）
        sample_pages = min(5, len(reader.pages))
        text_count = 0

        for i in range(sample_pages):
            page = reader.pages[i]
            text = page.extract_text() or ""
            # 去除空白后字数
            chars = len(text.strip())
            if chars > 50:
                text_count += 1

        scanned_pages = sample_pages - text_count
        result["scanned_pages"] = scanned_pages
        result["text_pages"] = text_count

        # 判断类型
        if scanned_pages == 0:
            result["type"] = "digital"
        elif scanned_pages == sample_pages:
            result["type"] = "scanned"
        else:
            result["type"] = "mixed"  # 混合：部分扫描+部分文字

        # 语言检测（基于提取的文字）
        all_text = ""
        for page in reader.pages[:sample_pages]:
            all_text += (page.extract_text() or "")

        # 简单的中英文检测
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', all_text))
        english_chars = len(re.findall(r'[a-zA-Z]', all_text))

        if chinese_chars > english_chars * 0.3 and chinese_chars > 20:
            result["lang"] = "zh"
        elif english_chars > 50:
            result["lang"] = "en"
        else:
            result["lang"] = "auto"

    except ImportError:
        print("[警告] pypdf 未安装，无法自动检测，使用默认配置（digital）", file=sys.stderr)
    except Exception as e:
        print(f"[警告] PDF 检测失败: {e}，使用默认配置（digital）", file=sys.stderr)

    return result


def start_hybrid_server(force_ocr: bool = False, ocr_lang: str = "auto",
                         enrich_formula: bool = False,
                         enrich_picture: bool = False,
                         port: int = HYBRID_SERVER_PORT) -> bool:
    """启动 Hybrid server（后台运行）"""

    # 检查是否已有运行中的 server
    if is_server_running(port):
        print(f"[INFO] Hybrid server 已在运行 (port {port})，跳过启动。")
        return True

    cmd = [JAVA_BIN, "-jar",
           "/home/wangyc/.local/lib/python3.10/site-packages/opendataloader_pdf/jar/opendataloader-pdf-hybrid.jar",
           "--port", str(port)]

    if force_ocr:
        cmd.append("--force-ocr")

    if ocr_lang and ocr_lang != "auto":
        cmd.extend(["--ocr-lang", ocr_lang])

    if enrich_formula:
        cmd.append("--enrich-formula")

    if enrich_picture:
        cmd.append("--enrich-picture-description")

    env = os.environ.copy()
    env["JAVA_HOME"] = JAVA_HOME
    env["PATH"] = JAVA_HOME + "/bin:" + env.get("PATH", "")

    print(f"[INFO] 启动 Hybrid server: {' '.join(cmd)}")

    try:
        proc = subprocess.Popen(
            cmd,
            env=env,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
        )

        # 等待 server 启动（最多30秒）
        for _ in range(30):
            time.sleep(1)
            if proc.poll() is not None:
                # 进程已退出
                output = proc.stdout.read().decode(errors="replace")
                print(f"[错误] Hybrid server 启动失败:\n{output[:500]}", file=sys.stderr)
                return False

            if is_server_running(port):
                # 写 PID 到文件
                with open(HYBRID_SERVER_PID_FILE, "w") as f:
                    f.write(str(proc.pid))
                print(f"[INFO] Hybrid server 启动成功 (PID: {proc.pid})")
                return True

        print("[错误] Hybrid server 启动超时", file=sys.stderr)
        proc.kill()
        return False

    except Exception as e:
        print(f"[错误] 启动 Hybrid server 失败: {e}", file=sys.stderr)
        return False


def is_server_running(port: int = HYBRID_SERVER_PORT) -> bool:
    """检查 server 是否在运行"""
    import socket
    try:
        with socket.create_connection(("127.0.0.1", port), timeout=2):
            return True
    except (socket.timeout, ConnectionRefusedError, OSError):
        return False


def stop_hybrid_server():
    """停止 Hybrid server"""
    pid_file = Path(HYBRID_SERVER_PID_FILE)
    if pid_file.exists():
        pid = int(pid_file.read_text().strip())
        try:
            os.kill(pid, 9)
            print(f"[INFO] Hybrid server 已停止 (PID: {pid})")
        except ProcessLookupError:
            print(f"[INFO] Hybrid server 进程已不存在")
        pid_file.unlink()
    else:
        # 尝试通过端口查找
        try:
            import socket
            s = socket.socket()
            s.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
            s.bind(("127.0.0.1", HYBRID_SERVER_PORT))
            s.close()
        except OSError:
            # port in use → try kill
            subprocess.run(["fuser", "-k", f"{HYBRID_SERVER_PORT}/tcp"],
                          capture_output=True)


def build_opendataloader_cmd(pdf_path: str, output_dir: str,
                              detection_info: dict,
                              output_format: str = "markdown,json") -> list:
    """根据检测结果构建 opendataloader-pdf 命令"""

    cmd = [
        sys.executable,  # 当前 Python 解释器
        "-m", "opendataloader_pdf",
        pdf_path,
        "-o", output_dir,
        "-f", output_format,
    ]

    pdf_type = detection_info["type"]

    if pdf_type == "digital":
        # 标准数字 PDF → 本地 Fast 模式，不需要 server
        return [sys.executable, "-m", "opendataloader_pdf",
                pdf_path, "-o", output_dir, "-f", output_format]

    else:
        # 扫描/mixed → Hybrid 模式
        hybrid_cmd = [
            sys.executable, "-m", "opendataloader_pdf",
            "--hybrid", "docling-fast",
        ]

        if pdf_type == "scanned" or detection_info["scanned_pages"] > 0:
            hybrid_cmd.append("--force-ocr")

        hybrid_cmd.extend([pdf_path, "-o", output_dir, "-f", output_format])

        return hybrid_cmd


def run_convert(pdf_path: str, output_dir: str,
                output_format: str = "markdown,json",
                auto_detect: bool = True,
                force_mode: str = None,
                skip_server: bool = False) -> dict:
    """
    主转换函数

    参数:
        pdf_path: PDF 文件路径
        output_dir: 输出目录
        output_format: 输出格式 (json, markdown, html, text 等)
        auto_detect: 是否自动检测 PDF 类型
        force_mode: 强制模式 ("fast" | "hybrid")
        skip_server: 跳过启动 server（server 已由外部启动）
    """

    result = {
        "success": False,
        "pdf_path": pdf_path,
        "output_dir": output_dir,
        "detection": None,
        "mode_used": None,
        "command": None,
        "stdout": "",
        "stderr": "",
        "returncode": None,
    }

    # ---- Excel 文档：openpyxl直接提取 + PDF位置辅助 ----
    ext_lower = pdf_path.lower()
    if ext_lower.endswith(".xlsx") or ext_lower.endswith(".xls"):
        print(f"[INFO] 检测到 Excel 文档 → openpyxl 提取 + PDF位置辅助")
        original_basename = Path(pdf_path).stem
        os.makedirs(output_dir, exist_ok=True)

        try:
            # Step 1: openpyxl 提取（结构化数据 + 表格 + 图表）
            print(f"[Step 1/3] Excel 数据提取...")
            excel_data = extract_excel_to_json(pdf_path)

            # Step 2: Excel → PDF（用于页码位置）
            print(f"[Step 2/3] Excel → PDF（用于提取位置信息）...")
            pdf_converted = convert_word_to_pdf(pdf_path, output_dir="/tmp")

            # Step 3: PDF 位置提取
            print(f"[Step 3/3] PDF 位置提取...")
            pdf_json_path = Path("/tmp") / f"{original_basename}.json"
            env = os.environ.copy()
            env["JAVA_HOME"] = JAVA_HOME
            env["PATH"] = JAVA_HOME + "/bin:" + env.get("PATH", "")
            subprocess.run(
                [sys.executable, "-m", "opendataloader_pdf",
                 pdf_converted, "-o", "/tmp", "-f", "json"],
                capture_output=True, text=True, env=env, timeout=300
            )

            # 匹配：按 sheet 名称匹配 page（改进版）
            pdf_page_by_sheet = {}
            pdf_page_count = 0

            if pdf_json_path.exists():
                with open(pdf_json_path, encoding="utf-8") as f:
                    pdf_data = json.load(f)

                pdf_page_count = pdf_data.get("number of pages", 0)

                # 策略1：收集每页的全部文本
                page_texts = {}  # page_num -> all_text
                for elem in pdf_data.get("flat_elements", []):
                    pg = elem.get("page number") or elem.get("page", "?")
                    txt = elem.get("content", "") or ""
                    if pg not in page_texts:
                        page_texts[pg] = ""
                    page_texts[pg] += " " + txt

                # 策略2：建立每页的标题词集合（用于模糊匹配）
                page_headings = {}  # page_num -> set of heading words
                for elem in pdf_data.get("flat_elements", []):
                    if elem.get("type") == "heading":
                        pg = elem.get("page number") or elem.get("page", "?")
                        if pg not in page_headings:
                            page_headings[pg] = set()
                        heading_text = elem.get("content", "") or ""
                        # 拆词
                        for w in re.findall(r"[\w\W]{2,}", heading_text):
                            page_headings[pg].add(w)

                # 为每个 sheet 匹配最可能的 page
                sheets = excel_data.get("sheets", [])

                for sh in sheets:
                    sh_normalized = sh.strip()

                    # 精确匹配：sheet 名出现在某页
                    matched = False
                    for pg, full_text in page_texts.items():
                        if sh_normalized and sh_normalized[:15] in full_text:
                            pdf_page_by_sheet[sh] = pg
                            matched = True
                            break

                    # 模糊匹配：共享词汇多的页面
                    if not matched:
                        sh_words = set(re.findall(r"[\w\W]{2,}", sh_normalized))
                        best_pg, best_score = None, 0
                        for pg, heading_words in page_headings.items():
                            if heading_words and sh_words:
                                score = len(sh_words & heading_words) / len(sh_words | heading_words)
                                if score > best_score and score > 0.1:
                                    best_score = score
                                    best_pg = pg
                        if best_pg:
                            pdf_page_by_sheet[sh] = best_pg
                            matched = True

                    # 未匹配到：不记录（页码留空，人工核对）
                    if not matched:
                        pdf_page_by_sheet[sh] = None

                pdf_json_path.unlink(missing_ok=True)

            # 为每个元素写入 PDF 页码（写入 location["pdf_page"]）
            matched_sheets = {}
            for elem in excel_data.get("elements", []):
                sh = elem.get("location", {}).get("excel", "") if isinstance(elem.get("location"), dict) else elem.get("sheet", "")
                if sh in pdf_page_by_sheet and pdf_page_by_sheet[sh]:
                    elem.setdefault("location", {})["pdf_page"] = pdf_page_by_sheet[sh]
                    matched_sheets[sh] = pdf_page_by_sheet[sh]
                else:
                    elem.setdefault("location", {})["pdf_page"] = None

            # 更新 excel_data 元数据
            excel_data["source"] = "openpyxl" + ("+pdf-position" if matched_sheets else "")
            excel_data["sheet_page_map"] = dict(matched_sheets)
            excel_data["pdf_total_pages"] = pdf_page_count
            excel_data["pdf_page_matched"] = len(matched_sheets)
            excel_data["pdf_page_unmatched"] = len(sheets) - len(matched_sheets)

            # 输出 sheet→页码映射供核对（仅显示已匹配到的）
            matched_sheets = {sh: pg for sh, pg in pdf_page_by_sheet.items() if pg}
            if matched_sheets:
                print(f"[INFO] Sheet → PDF页码 映射（共 {pdf_page_count} 页，已匹配 {len(matched_sheets)}/{len(sheets)} 个Sheet）:")
                for sh, pg in matched_sheets.items():
                    print(f"       Sheet「{sh}」→ PDF第 {pg} 页")
            if len(matched_sheets) < len(sheets):
                unmatched = [sh for sh in sheets if sh not in matched_sheets]
                print(f"[INFO] 以下 Sheet 未匹配到页码（PDF中未出现Sheet名，请人工核对）:")
                for sh in unmatched:
                    print(f"       Sheet「{sh}」→ 需人工确认页码")


            # 保存 JSON
            json_path = Path(output_dir) / f"{original_basename}.json"
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(excel_data, f, ensure_ascii=False, indent=2)

            # 生成 Markdown
            md_content = extract_excel_to_markdown(pdf_path, page_hints=pdf_page_by_sheet)
            md_path = Path(output_dir) / f"{original_basename}.md"
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(md_content)

            # 清理临时 PDF
            Path(pdf_converted).unlink(missing_ok=True)

            print(f"[成功] Excel 提取完成！")
            print(f"[INFO] 生成文件:")
            print(f"  JSON: {json_path}")
            print(f"  MD:   {md_path}")
            print(f"[INFO] 包含: {excel_data['total_sheets']} 个Sheet | "
                  f"{excel_data['total_tables']} 张表格 | "
                  f"{excel_data['total_charts']} 个图表")

            return {
                "success": True,
                "output_dir": output_dir,
                "mode_used": "excel-openpyxl",
                "files_created": [str(json_path), str(md_path)],
                "excel_stats": {
                    "sheets": excel_data["total_sheets"],
                    "tables": excel_data["total_tables"],
                    "charts": excel_data["total_charts"],
                }
            }

        except Exception as e:
            print(f"[错误] Excel 提取失败: {e}")
            import traceback
            traceback.print_exc()
            return {"success": False, "mode_used": "excel-failed", "error": str(e)}

    # ---- Word 文档：docx内容 + PDF位置 → 合并 ----
    ext_lower = pdf_path.lower()
    if ext_lower.endswith(".docx") or ext_lower.endswith(".doc"):
        print(f"[INFO] 检测到 Word 文档 → 三步处理：docx提取 + PDF位置提取 → 智能合并")
        original_basename = Path(pdf_path).stem
        os.makedirs(output_dir, exist_ok=True)

        env = os.environ.copy()
        env["JAVA_HOME"] = JAVA_HOME
        env["PATH"] = JAVA_HOME + "/bin:" + env.get("PATH", "")

        try:
            # Step 1: Word → PDF
            print(f"[Step 1/3] Word → PDF（用于提取位置信息）...")
            pdf_converted = convert_word_to_pdf(pdf_path, output_dir="/tmp")

            # Step 2: PDF 位置提取（page / bbox / section_path）
            print(f"[Step 2/3] PDF 位置提取...")
            pdf_json_cmd = [
                sys.executable, "-m", "opendataloader_pdf",
                pdf_converted, "-o", "/tmp", "-f", "json",
            ]
            subprocess.run(pdf_json_cmd, capture_output=True, text=True, env=env, timeout=300)

            # opendataloader 输出的文件名就是 basename.json
            pdf_json_path = Path("/tmp") / f"{original_basename}.json"
            pdf_flat = []
            if pdf_json_path.exists():
                with open(pdf_json_path, encoding="utf-8") as f:
                    pdf_data = json.load(f)
                pdf_flat = pdf_data.get("flat_elements", [])
                # 兼容：无 flat_elements 时用 kids 扁平化
                if not pdf_flat and "kids" in pdf_data:
                    pdf_flat = []
                    heading_path = []
                    def _flat(kids, out, hp):
                        for k in kids:
                            if k.get("type") == "heading":
                                level = k.get("heading level", 99)
                                while hp and hp[-1][0] >= level:
                                    hp.pop()
                                hp.append((level, k.get("content", "")))
                                k["section_path"] = " > ".join(x[1] for x in sorted(hp))
                                out.append(k)
                                if "kids" in k:
                                    _flat(k.pop("kids"), out, hp)
                            else:
                                out.append(k)
                                if "kids" in k:
                                    _flat(k.pop("kids"), out, hp)
                    _flat(pdf_data["kids"], pdf_flat, heading_path)
                pdf_json_path.unlink(missing_ok=True)  # 清理临时文件

            print(f"[      ] PDF 提取到 {len(pdf_flat)} 个位置元素")

            # Step 3: docx 内容提取（表格完整 + 文本完整）
            print(f"[Step 3/3] docx 内容提取（完整表格/文本）...")
            docx_data = extract_word_to_json(pdf_path)  # 用已有的 docx 提取函数

            # Step 4: 智能合并
            print(f"[      ] 智能合并：docx内容 + PDF位置 → 统一结构")
            merged = merge_docx_and_pdf(docx_data, pdf_flat)

            # 附加元数据
            merged["source"] = "word-to-pdf-merged"
            merged["original_word"] = os.path.abspath(pdf_path)
            merged["converted_pdf"] = pdf_converted

            # 保存合并后的 JSON
            merged_json_path = Path(output_dir) / f"{original_basename}.json"
            with open(merged_json_path, "w", encoding="utf-8") as f:
                json.dump(merged, f, ensure_ascii=False, indent=2)

            # 生成 Markdown（基于 docx 内容，带 section_path）
            md_content = f"# {original_basename}\n\n"
            md_content += convert_word_to_markdown(pdf_path)
            md_path = Path(output_dir) / f"{original_basename}.md"
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(md_content)

            match_stats = merged.get("_pdf_position_hints", {})
            print(f"[成功] 合并完成！")
            print(f"[INFO] 生成文件:")
            print(f"  JSON: {merged_json_path}  (内容:{len(merged['elements'])}元素 | 表格:{merged['total_tables']} | 段落:{merged['total_paragraphs']})")
            print(f"  PDF位置匹配: 表格匹配{match_stats.get('tables_matched','?')}个, 段落匹配{match_stats.get('paras_matched','?')}个")
            print(f"  MD:   {md_path}")
            print(f"[INFO] 临时PDF已清理（{pdf_converted}）")

            # 清理临时 PDF
            Path(pdf_converted).unlink(missing_ok=True)

            return {
                "success": True,
                "original_word": os.path.abspath(pdf_path),
                "output_dir": output_dir,
                "mode_used": "word-to-pdf-merged",
                "files_created": [str(merged_json_path), str(md_path)],
                "merged_stats": match_stats,
            }

        except Exception as e:
            print(f"[错误] 合并流程失败: {e}")
            import traceback
            traceback.print_exc()
            # 降级：仅 docx 直接提取
            word_result = convert_word(pdf_path, output_dir, output_format)
            word_result["mode_used"] = "word-direct-fallback"
            return word_result
        word_result = convert_word(pdf_path, output_dir, output_format)
        word_result["mode_used"] = "word-direct-fallback"
        return word_result

    # ---- PDF 处理 ----
    # pdf_path 已在参数中传入，继续使用

    # 验证 Java
    if not find_java():
        print("[错误] Java 不可用，请检查 JAVA_HOME 设置", file=sys.stderr)
        return result

    # 解析 PDF 路径
    pdf_path = os.path.abspath(pdf_path)
    if not os.path.exists(pdf_path):
        print(f"[错误] 文件不存在: {pdf_path}", file=sys.stderr)
        return result

    # 自动检测
    if auto_detect and not force_mode:
        print("[INFO] 自动检测 PDF 类型...")
        detection = detect_pdf_type(pdf_path)
        result["detection"] = detection
        print(f"[INFO] 检测结果: {json.dumps(detection, ensure_ascii=False, indent=2)}")
    else:
        detection = {
            "type": "digital" if (force_mode == "fast" or not force_mode) else "scanned",
            "lang": "auto",
            "total_pages": 0,
        }
        result["detection"] = detection

    # 构建命令
    cmd = build_opendataloader_cmd(pdf_path, output_dir, detection, output_format)
    result["command"] = " ".join(cmd)

    print(f"\n[INFO] 使用模式: {'Fast (本地)' if detection['type'] == 'digital' else 'Hybrid (AI增强)'}")
    print(f"[INFO] 执行命令: {' '.join(cmd)}")

    # 启动 Hybrid server（如果需要且未运行）
    if detection["type"] != "digital" and not skip_server:
        # 准备 server 参数
        force_ocr = detection["type"] in ("scanned", "mixed")
        ocr_lang = detection.get("lang", "auto")
        if ocr_lang == "auto":
            ocr_lang = "zh,en"  # 默认中英文

        enrich_formula = detection.get("has_formulas", False)
        enrich_picture = detection.get("has_charts", False)

        server_ok = start_hybrid_server(
            force_ocr=force_ocr,
            ocr_lang=ocr_lang,
            enrich_formula=enrich_formula,
            enrich_picture=enrich_picture
        )

        if not server_ok:
            print("[错误] Hybrid server 启动失败，尝试降级到 Fast 模式...", file=sys.stderr)
            detection["type"] = "digital"
            cmd = build_opendataloader_cmd(pdf_path, output_dir, detection, output_format)
            result["command"] = " ".join(cmd)

    # 执行
    env = os.environ.copy()
    env["JAVA_HOME"] = JAVA_HOME
    env["PATH"] = JAVA_HOME + "/bin:" + env.get("PATH", "")

    try:
        proc = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            env=env,
            timeout=300  # 5分钟超时
        )
        result["stdout"] = proc.stdout
        result["stderr"] = proc.stderr
        result["returncode"] = proc.returncode

        if proc.returncode == 0:
            result["success"] = True
            result["mode_used"] = detection["type"]
            print(f"[成功] 输出已保存到: {output_dir}")
        else:
            print(f"[错误] 转换失败 (返回码 {proc.returncode}):", file=sys.stderr)
            print(proc.stderr[:500] if proc.stderr else proc.stdout[:500], file=sys.stderr)

    except subprocess.TimeoutExpired:
        result["stderr"] = "超时（5分钟）"
        print("[错误] 转换超时（5分钟）", file=sys.stderr)
    except Exception as e:
        result["stderr"] = str(e)
        print(f"[错误] {e}", file=sys.stderr)

    return result


# ---------- CLI 入口 ----------
def main():
    parser = argparse.ArgumentParser(
        description="PDF/Word 文档自动检测与转换工具（统一接口）"
    )
    parser.add_argument("input", help="输入 PDF 文件路径")
    parser.add_argument("-o", "--output", required=True, help="输出目录")
    parser.add_argument("-f", "--format",
                        default="markdown,json",
                        help="输出格式（逗号分隔，默认 markdown,json）")
    parser.add_argument("--no-auto-detect", dest="auto_detect",
                        action="store_false", default=True,
                        help="禁用自动检测")
    parser.add_argument("--force-mode",
                        choices=["fast", "hybrid"],
                        default=None,
                        help="强制指定模式")
    parser.add_argument("--skip-server", action="store_true",
                        help="跳过 server 启动（server 已运行时使用）")
    parser.add_argument("--stop-server", action="store_true",
                        help="停止 Hybrid server")
    parser.add_argument("--detect-only", action="store_true",
                        help="仅检测 PDF 类型，不执行转换")

    args = parser.parse_args()

    # 处理 stop-server
    if args.stop_server:
        stop_hybrid_server()
        return

    # 处理 detect-only
    if args.detect_only:
        ext = args.input.lower()
        if ext.endswith(".docx") or ext.endswith(".doc"):
            info = {"type": "word", "format": "docx/doc"}
        elif ext.endswith(".xlsx") or ext.endswith(".xls"):
            info = {"type": "excel", "format": "xlsx/xls"}
        else:
            info = detect_pdf_type(args.input)
        print(json.dumps(info, ensure_ascii=False, indent=2))
        return

    # 创建输出目录
    os.makedirs(args.output, exist_ok=True)

    # 执行转换
    result = run_convert(
        pdf_path=args.input,
        output_dir=args.output,
        output_format=args.format,
        auto_detect=args.auto_detect,
        force_mode=args.force_mode,
        skip_server=args.skip_server,
    )

    # 输出结果 JSON
    print("\n--- Result ---")
    print(json.dumps({
        "success": result["success"],
        "mode": result["mode_used"],
        "detection": result.get("detection"),
        "files_created": result.get("files_created", []),
        "command": result.get("command"),
    }, ensure_ascii=False, indent=2))

    sys.exit(0 if result["success"] else 1)


if __name__ == "__main__":
    main()
